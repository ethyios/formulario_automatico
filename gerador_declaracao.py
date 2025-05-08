import tkinter as tk
from tkinter import ttk, messagebox
import os
import datetime
import sys
import subprocess
from docx import Document
from docx2pdf import convert
import io
import importlib
import importlib.util # Adicionado para importação dinâmica

# Redirect stdout/stderr if they are None (common in PyInstaller --noconsole apps)
# This should be one of the very first things the application does.
if sys.stdout is None:
    sys.stdout = open(os.devnull, 'w')
if sys.stderr is None:
    sys.stderr = open(os.devnull, 'w')

# --- Constantes para o modelo ---
HIDDEN_FOLDER_NAME = "._modelo_data"
MODEL_BASENAME = "declaracao_base_bytes.py"

# --- Importar função da GUI do importador e carregar bytes iniciais ---
DOCX_BYTES = None # Inicializa como None

def get_application_path():
    """Retorna o caminho base da aplicação, seja script ou executável."""
    if getattr(sys, 'frozen', False): # Rodando como bundle PyInstaller
        return os.path.dirname(sys.executable)
    else: # Rodando como script
        return os.path.dirname(os.path.abspath(__file__))

def _load_module_from_path(module_name, file_path):
    """Carrega um módulo dinamicamente a partir de um caminho de arquivo."""
    spec = importlib.util.spec_from_file_location(module_name, file_path)
    if spec is None:
        raise ImportError(f"Não foi possível encontrar o spec para {module_name} em {file_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module # Adiciona ao sys.modules ANTES de executar
    spec.loader.exec_module(module)
    return module

def carregar_docx_bytes_inicialmente():
    global DOCX_BYTES
    app_path = get_application_path()
    model_dir_path = os.path.join(app_path, HIDDEN_FOLDER_NAME)
    model_file_path = os.path.join(model_dir_path, MODEL_BASENAME)
    module_name_to_load = "declaracao_base_bytes_dinamico" # Nome único para evitar conflitos

    try:
        # Garante que o diretório do modelo esteja no sys.path para importação direta se necessário
        # ou para que o importlib.util.spec_from_file_location funcione corretamente.
        # No entanto, para PyInstaller, é mais robusto carregar explicitamente pelo caminho.
        
        if not os.path.exists(model_file_path):
            DOCX_BYTES = None
            # print(f"DEBUG: Arquivo do modelo '{model_file_path}' não encontrado na carga inicial.")
            return

        # Remove o módulo antigo se existir para forçar o recarregamento do arquivo
        if module_name_to_load in sys.modules:
            del sys.modules[module_name_to_load]
        
        reloaded_module = _load_module_from_path(module_name_to_load, model_file_path)
        DOCX_BYTES = reloaded_module.DOCX_BYTES
    except ImportError:
        DOCX_BYTES = None
        # print(f"DEBUG: Falha ao importar '{MODEL_BASENAME}' de '{model_dir_path}' na carga inicial.")
    except AttributeError:
        DOCX_BYTES = None
        # print(f"DEBUG: 'DOCX_BYTES' não encontrado em '{MODEL_BASENAME}' na carga inicial.")
    except Exception as e:
        DOCX_BYTES = None
        # print(f"DEBUG: Erro inesperado ao carregar DOCX_BYTES inicialmente: {e}")

carregar_docx_bytes_inicialmente() # Tenta carregar ao iniciar o script

try:
    from importar_declaracao import iniciar_interface_importador
except ImportError:
    def iniciar_interface_importador(): # Fallback se o import falhar
        messagebox.showerror("Erro Crítico",
                             "Não foi possível encontrar o módulo 'importar_declaracao.py'.\n"
                             "A funcionalidade de importação de modelo não está disponível.")
    # print("DEBUG: Falha ao importar 'iniciar_interface_importador' de 'importar_declaracao'.")


# --- Configuração ---
OUTPUT_FOLDER_NAME = "declaracoes_geradas"
PLACEHOLDERS = {
    "nome_responsavel": "{{NOME_RESPONSAVEL}}",
    "nome_filho": "{{NOME_FILHO}}",
    "serie": "{{SERIE}}",
    "data": "{{DATA}}",
    "periodo": "{{PERIODO}}"  # Novo placeholder
}

def formatar_data_por_extenso(data_str):
    """Converte uma data de 'DD/MM/AAAA' para 'DD de [mês] de AAAA'."""
    try:
        dia, mes_num, ano = data_str.split('/')
        meses = {
            "01": "janeiro",
            "02": "fevereiro",
            "03": "março",
            "04": "abril",
            "05": "maio", 
            "06": "junho",
            "07": "julho",
            "08": "agosto",
            "09": "setembro",
            "10": "outubro",
            "11": "novembro",
            "12": "dezembro"
        }
        mes_extenso = meses.get(mes_num, "")
        if not mes_extenso:
            return data_str # Retorna original se o mês for inválido
        return f"{dia} de {mes_extenso} de {ano}"
    except ValueError:
        # Retorna a string original se o formato for inesperado
        return data_str

def apply_replacements(doc, replacements):
    """
    Substitui os placeholders no documento (parágrafos e tabelas).
    'replacements' é um dicionário como {'{{PLACEHOLDER}}': 'Valor Real'}
    """
    # Substitui nos parágrafos do corpo principal
    for p in doc.paragraphs:
        original_text = p.text
        modified_text = original_text
        for placeholder, value in replacements.items():
            # Garante que o valor da substituição seja uma string
            modified_text = modified_text.replace(placeholder, str(value))
        
        if original_text != modified_text:
            # Atribuir a p.text substitui todo o conteúdo do parágrafo.
            # Isso lida com placeholders divididos em várias 'runs',
            # mas pode simplificar a formatação de 'run' dentro do parágrafo.
            p.text = modified_text

    # Substitui nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p_cell in cell.paragraphs:
                    original_cell_text = p_cell.text
                    modified_cell_text = original_cell_text
                    for placeholder, value in replacements.items():
                        modified_cell_text = modified_cell_text.replace(placeholder, str(value))
                    
                    if original_cell_text != modified_cell_text:
                        p_cell.text = modified_cell_text

def recarregar_modelo_docx():
    global DOCX_BYTES
    app_path = get_application_path()
    model_dir_path = os.path.join(app_path, HIDDEN_FOLDER_NAME)
    model_file_path = os.path.join(model_dir_path, MODEL_BASENAME)
    module_name_to_load = "declaracao_base_bytes_dinamico" # Mesmo nome usado na carga inicial

    try:
        if not os.path.exists(model_file_path):
            DOCX_BYTES = None
            messagebox.showerror("Erro ao Recarregar",
                                 f"Arquivo '{MODEL_BASENAME}' não encontrado em '{model_dir_path}'.\n"
                                 "Use o botão 'Importar/Atualizar Modelo DOCX'.")
            status_label.config(text=f"Erro: '{MODEL_BASENAME}' não encontrado.")
            app.update_idletasks()
            return False

        # Remove o módulo antigo se existir para forçar o recarregamento do arquivo
        if module_name_to_load in sys.modules:
            del sys.modules[module_name_to_load]

        reloaded_module = _load_module_from_path(module_name_to_load, model_file_path)
        DOCX_BYTES = reloaded_module.DOCX_BYTES
        messagebox.showinfo("Sucesso", f"O modelo DOCX foi recarregado com sucesso de '{model_file_path}'.")
        status_label.config(text="Modelo DOCX recarregado com os novos bytes.")
        app.update_idletasks()
        return True
    except ImportError:
        DOCX_BYTES = None
        messagebox.showerror("Erro ao Recarregar",
                             f"Não foi possível importar '{MODEL_BASENAME}' de '{model_dir_path}'.\n"
                             "Use o botão 'Importar/Atualizar Modelo DOCX'.")
        status_label.config(text=f"Erro: Falha ao importar '{MODEL_BASENAME}'.")
        app.update_idletasks()
        return False
    except AttributeError: # Se DOCX_BYTES não estiver no módulo recarregado
        DOCX_BYTES = None
        messagebox.showerror("Erro ao Recarregar",
                             f"'DOCX_BYTES' não encontrado no arquivo '{MODEL_BASENAME}' (em '{model_dir_path}').\n"
                             "Verifique se o modelo foi importado corretamente ou importe novamente.")
        status_label.config(text="Erro: 'DOCX_BYTES' ausente no modelo importado.")
        app.update_idletasks()
        return False
    except Exception as e:
        DOCX_BYTES = None
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro inesperado ao recarregar o modelo: {e}")
        status_label.config(text=f"Erro ao recarregar modelo: {e}")
        app.update_idletasks()
        return False

def abrir_janela_importador_e_recarregar():
    status_label.config(text="Abrindo interface para importação de modelo...")
    app.update_idletasks()

    # Chama a interface do importador (que é modal)
    iniciar_interface_importador()

    # Após o fechamento da janela do importador, tenta recarregar o modelo
    status_label.config(text="Tentando recarregar o modelo DOCX...")
    app.update_idletasks()
    recarregar_modelo_docx()
    # A função recarregar_modelo_docx já atualiza o status_label e mostra popups.

def gerar_declaracao():
    nome_responsavel = entry_nome_responsavel.get()
    nome_filho = entry_nome_filho.get()
    serie = entry_serie.get()
    data_declaracao = entry_data.get()
    periodo = entry_periodo.get() # Novo campo

    if not all([nome_responsavel, nome_filho, serie, data_declaracao, periodo]): # Adicionado periodo à validação
        messagebox.showerror("Erro de Validação", "Todos os campos são obrigatórios!")
        return

    progress_bar['value'] = 0
    status_label.config(text="Iniciando processo...")
    app.update_idletasks()

    try:
        # 1. Verifica se o conteúdo do template foi importado
        if DOCX_BYTES is None:
            messagebox.showerror("Erro de Configuração",
                                 f"Os dados do modelo DOCX não foram encontrados (deveriam estar em '{os.path.join(HIDDEN_FOLDER_NAME, MODEL_BASENAME)}').\n"
                                 "Clique no botão 'Importar/Atualizar Modelo DOCX' para selecionar o arquivo .docx modelo e gerar/atualizar os dados necessários.")
            status_label.config(text="Erro: Modelo DOCX não carregado. Use o botão de importação.")
            return

        # Carrega o documento a partir dos bytes em memória
        try:
            doc_stream = io.BytesIO(DOCX_BYTES)
            doc = Document(doc_stream)
        except Exception as e_load:
            messagebox.showerror("Erro ao Carregar Modelo",
                                 f"Não foi possível carregar o modelo DOCX a partir dos dados importados: {e_load}\n"
                                 f"Verifique se o arquivo '{MODEL_BASENAME}' (em '{HIDDEN_FOLDER_NAME}') foi gerado corretamente.")
            status_label.config(text="Erro: Falha ao carregar modelo importado.")
            return
            
        progress_bar['value'] = 10
        status_label.config(text="Modelo carregado...")
        app.update_idletasks()

        # 2. Prepara e aplica as substituições
        data_formatada = formatar_data_por_extenso(data_declaracao)
        replacements_dict = {
            PLACEHOLDERS["nome_responsavel"]: nome_responsavel,
            PLACEHOLDERS["nome_filho"]: nome_filho,
            PLACEHOLDERS["serie"]: serie,
            PLACEHOLDERS["data"]: data_formatada,
            PLACEHOLDERS["periodo"]: periodo # Adicionado periodo ao dicionário
        }
        apply_replacements(doc, replacements_dict)
        progress_bar['value'] = 40
        status_label.config(text="Dados preenchidos no modelo...")
        app.update_idletasks()

        # 3. Cria a pasta de saída se não existir
        if not os.path.exists(OUTPUT_FOLDER_NAME):
            os.makedirs(OUTPUT_FOLDER_NAME)

        # Sanitiza componentes do nome do arquivo para evitar erros
        safe_nome_filho = "".join(c if c.isalnum() else "_" for c in nome_filho)
        safe_data = "".join(c if c.isalnum() else "_" for c in data_declaracao.replace("/", "-"))

        # Define nomes para os arquivos temporário e final
        temp_docx_filename = f"temp_declaracao_{safe_nome_filho}_{safe_data}.docx"
        output_pdf_filename = f"Declaracao_{safe_nome_filho}_{safe_data}.pdf"
        
        temp_docx_path = os.path.join(OUTPUT_FOLDER_NAME, temp_docx_filename)
        output_pdf_path = os.path.join(OUTPUT_FOLDER_NAME, output_pdf_filename)

        # 4. Salva o DOCX preenchido temporariamente
        doc.save(temp_docx_path)
        progress_bar['value'] = 60
        status_label.config(text="Documento Word temporário salvo...")
        app.update_idletasks()

        # 5. Converte para PDF
        status_label.config(text="Convertendo para PDF...")
        app.update_idletasks()
        convert(temp_docx_path, output_pdf_path)
        progress_bar['value'] = 90
        status_label.config(text="PDF gerado...")
        app.update_idletasks()

        # 6. Limpa o DOCX temporário
        os.remove(temp_docx_path)
        progress_bar['value'] = 100
        status_label.config(text="Declaração gerada com sucesso!")
        app.update_idletasks()

        messagebox.showinfo("Sucesso", f"Declaração gerada com sucesso!\nSalvo como: {output_pdf_path}")
        
        # 7. Abre a pasta de saída
        output_dir_abs_path = os.path.abspath(OUTPUT_FOLDER_NAME)
        try:
            if sys.platform == "win32": # Para Windows
                os.startfile(output_dir_abs_path)
            elif sys.platform == "darwin": # Para macOS
                 subprocess.run(['open', output_dir_abs_path], check=True)
            else: # Para Linux e outros Unix-like (assumindo xdg-utils)
                 subprocess.run(['xdg-open', output_dir_abs_path], check=True)
        except Exception as e_open:
            messagebox.showwarning("Aviso", "A declaração foi gerada, mas não foi possível abrir a pasta automaticamente.\n"
                                           f"Você pode encontrá-la em: {output_dir_abs_path}")

    except Exception as e:
        progress_bar['value'] = 0
        status_label.config(text="Erro ao gerar declaração.")
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro durante a geração:\n{e}")

# --- Configuração da Interface Gráfica (GUI) ---
app = tk.Tk()
app.title("Gerador de Declaração de Comparecimento")

# Estilo ttk
style = ttk.Style(app)
if "clam" in style.theme_names(): # Tenta usar um tema mais moderno se disponível
    style.theme_use('clam')

frame = ttk.Frame(app, padding="20")
frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
app.columnconfigure(0, weight=1)
app.rowconfigure(0, weight=1)

# Labels e Campos de Entrada
ttk.Label(frame, text="Nome completo do(a) responsável:").grid(row=0, column=0, sticky=tk.W, pady=(0,5))
entry_nome_responsavel = ttk.Entry(frame, width=50)
entry_nome_responsavel.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0,10))

ttk.Label(frame, text="Nome completo do filho(a):").grid(row=2, column=0, sticky=tk.W, pady=(0,5))
entry_nome_filho = ttk.Entry(frame, width=50)
entry_nome_filho.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(0,10))

ttk.Label(frame, text="Série:").grid(row=4, column=0, sticky=tk.W, pady=(0,5))
entry_serie = ttk.Entry(frame, width=50)
entry_serie.grid(row=5, column=0, sticky=(tk.W, tk.E), pady=(0,10))

ttk.Label(frame, text="Data da declaração (DD/MM/AAAA):").grid(row=6, column=0, sticky=tk.W, pady=(0,5))
entry_data = ttk.Entry(frame, width=50)
entry_data.grid(row=7, column=0, sticky=(tk.W, tk.E), pady=(0,10))
entry_data.insert(0, datetime.date.today().strftime("%d/%m/%Y")) # Preenche com data atual

# Novo campo para Período
ttk.Label(frame, text="Período (ex: matutino, vespertino, etc.):").grid(row=8, column=0, sticky=tk.W, pady=(0,5))
entry_periodo = ttk.Entry(frame, width=50)
entry_periodo.grid(row=9, column=0, sticky=(tk.W, tk.E), pady=(0,10))

# Botão para Importar/Atualizar Modelo
btn_importar = ttk.Button(frame, text="Importar/Atualizar Modelo DOCX", command=abrir_janela_importador_e_recarregar)
btn_importar.grid(row=10, column=0, sticky=(tk.W, tk.E), pady=(10,5)) # Ajustado row index

# Botão Gerar Declaração
btn_gerar = ttk.Button(frame, text="Gerar Declaração em PDF", command=gerar_declaracao)
btn_gerar.grid(row=11, column=0, sticky=(tk.W, tk.E), pady=(5,10)) # Ajustado row index

# Barra de Progresso
progress_bar = ttk.Progressbar(frame, orient="horizontal", length=300, mode="determinate")
progress_bar.grid(row=12, column=0, sticky=(tk.W, tk.E), pady=(0,5)) # Ajustado row index

# Label de Status
status_label = ttk.Label(frame, text="Preencha os campos e clique em gerar.")
status_label.grid(row=13, column=0, sticky=(tk.W, tk.E), pady=(0,0)) # Ajustado row index

# Configura o redimensionamento da coluna no frame
frame.columnconfigure(0, weight=1)

# Função para garantir o encerramento completo da aplicação
def on_closing():
    try:
        app.destroy()
    except tk.TclError:
        # Pode acontecer se a app já estiver sendo destruída
        pass
    sys.exit(0) # Garante que o processo Python termine

app.protocol("WM_DELETE_WINDOW", on_closing) # Intercepta o fechamento da janela
app.mainloop()